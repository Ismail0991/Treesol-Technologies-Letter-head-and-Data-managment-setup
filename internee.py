from flask import Flask, render_template, request, redirect, url_for, flash, send_file, session 
from google.cloud import firestore
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert
import os

app = Flask(__name__)
app.secret_key = "your_secret_key"

# Firestore client
db = firestore.Client.from_service_account_json("traineedata-a1379-8c9c23dd84c8.json")

# -------------------------
# Login System
# -------------------------
USERNAME = "rizwan89"
PASSWORD = "1234567891"

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]

        if username == USERNAME and password == PASSWORD:
            session["user"] = username
            flash("✅ Logged in successfully!", "success")
            return redirect(url_for("index"))
        else:
            flash("❌ Invalid username or password", "danger")

    return render_template("login.html")

@app.route("/logout")
def logout():
    session.pop("user", None)
    return redirect(url_for("login"))

# -------------------------
# Protect main routes
# -------------------------
@app.before_request
def require_login():
    allowed_routes = ["login", "static"]
    if "user" not in session and request.endpoint not in allowed_routes:
        return redirect(url_for("login"))

# -------------------------
# Home → Show internee list
# -------------------------
@app.route("/")
def index():
    docs = db.collection("internees").stream()
    internees = []
    today = datetime.today().date()

    is_direct_open = request.referrer is None or request.referrer.endswith(request.host_url)

    if is_direct_open:
        for doc in docs:
            d = doc.to_dict()
            d["id"] = doc.id
            try:
                end_date = datetime.strptime(d["end"], "%Y-%m-%d").date()
                days_left = (end_date - today).days
                if days_left in [2, 3]:
                    flash(f"⚠️ {d['name']}'s internship ends in {days_left} days!", "warning")
            except Exception as e:
                print("Date parse error:", e)
            internees.append(d)
    else:
        for doc in docs:
            d = doc.to_dict()
            d["id"] = doc.id
            internees.append(d)

    return render_template("index.html", internees=internees)


# Add internee
@app.route("/add", methods=["POST"])
def add_internee():
    data = {
        "name": request.form["name"],
        "father": request.form["father"],
        "cnic": request.form["cnic"],
        "phone": request.form["phone"],
        "field": request.form["field"],
        "start": request.form["start"],
        "end": request.form["end"],
    }
    db.collection("internees").add(data)
    flash("✅ Internee Added Successfully!", "success")
    return redirect(url_for("index"))


# Edit internee
@app.route("/edit/<id>", methods=["GET", "POST"])
def edit_internee(id):
    doc_ref = db.collection("internees").document(id)
    data = doc_ref.get().to_dict()

    if request.method == "POST":
        doc_ref.update({
            "name": request.form["name"],
            "father": request.form["father"],
            "cnic": request.form["cnic"],
            "phone": request.form["phone"],
            "field": request.form["field"],
            "start": request.form["start"],
            "end": request.form["end"]
        })
        flash("✅ Internee Updated Successfully!", "success")
        return redirect(url_for("index"))

    return render_template("edit.html", internee=data, id=id)


# Delete internee
@app.route("/delete/<id>")
def delete_internee(id):
    db.collection("internees").document(id).delete()
    flash("❌ Internee Deleted Successfully!", "danger")
    return redirect(url_for("index"))


# Generate internship completion letter (PDF)
@app.route("/letter/<id>", methods=["POST"])
def generate_letter(id):
    internee = db.collection("internees").document(id).get().to_dict()
    if not internee:
        flash("❌ Internee not found!", "danger")
        return redirect(url_for("index"))

    start = internee["start"]
    end = internee["end"]

    # Create Word document
    doc = Document()

    # Letterhead: logo only (centered at top)
    if os.path.exists("static/s1.png"):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r = p.add_run()
        r.add_picture("static/s1.png", width=Inches(5.0))

    doc.add_paragraph("")  # spacing

    # Body text (justified alignment, professional style)
    body = doc.add_paragraph()
    body.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Bold intern's name and field
    run = body.add_run(
        f"We are pleased to confirm that "
    )
    run.font.size = Pt(12)

    run_name = body.add_run(f"{internee['name']} ")
    run_name.bold = True
    run_name.font.size = Pt(12)

    run2 = body.add_run(f"Son/Daughter of {internee['father']} worked as a ")
    run2.font.size = Pt(12)

    run_field = body.add_run(f"{internee['field']} Intern ")
    run_field.bold = True
    run_field.font.size = Pt(12)

    run3 = body.add_run(
        f"in our firm TreeSol Technologies PVT Ltd. "
        f"He/She has completed his/her internship from {start} to {end}. "
        f"During the internship, he/she demonstrated good {internee['field']} skills with a self-motivated attitude "
        f"to learn new things. We wish him/her all the best for his/her future endeavours.\n"
        f"Warm Regards"
    )
    run3.font.size = Pt(12)

    
    # 
# Issued date (one line above Warm Regards)
    doc.add_paragraph("Issued on: " + datetime.today().strftime("%d-%m-%Y"))

    # Stamp image aligned right
    if os.path.exists("static/stamp.png"):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        r = p.add_run()
        r.add_picture("static/stamp.png", width=Inches(1.2))

    # ✅ Add s2.png at the bottom (footer, centered)
    if os.path.exists("static/s2.png"):
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r = p.add_run()
        r.add_picture("static/s2.png", width=Inches(6.5))

    # Save DOCX first
    filename_docx = f"internship_letter_{internee['name'].replace(' ', '_')}.docx"
    filepath_docx = os.path.join("letters", filename_docx)
    os.makedirs("letters", exist_ok=True)
    doc.save(filepath_docx)

    # Convert to PDF
    filepath_pdf = filepath_docx.replace(".docx", ".pdf")
    convert(filepath_docx, filepath_pdf)

    return send_file(filepath_pdf, as_attachment=True)


from waitress import serve
if __name__ == "__main__":
 serve(app, host="0.0.0.0", port=8080)
